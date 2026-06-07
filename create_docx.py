import os
import re
from docx import Document
from docx.shared import Pt, Inches

def md_to_docx(md_path, docx_template_path, output_docx_path):
    print("Loading Template...")
    # Load template document to keep headers, footers, styles, page layout
    doc = Document(docx_template_path)
    
    print("Clearing paragraphs...")
    # Remove all existing paragraphs so we start fresh but with same structure
    for p in doc.paragraphs:
        p_element = p._p
        parent_element = p_element.getparent()
        if parent_element is not None:
            parent_element.remove(p_element)
            
    print("Reading Markdown...")
    # Read Markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
        
    lines = md_text.split('\n')
    
    in_code_block = False
    
    for line in lines:
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            # Apply code styling
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue
            
        line_clean = line.strip()
        
        # Horizontal Rule
        if line_clean == '---':
            doc.add_paragraph()
            continue
            
        if not line_clean:
            doc.add_paragraph()
            continue
            
        if line_clean.startswith('# '):
            text = line_clean[2:].strip()
            try:
                doc.add_heading(text, level=1)
            except KeyError:
                p = doc.add_paragraph(text)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(16)
        elif line_clean.startswith('## '):
            text = line_clean[3:].strip()
            try:
                doc.add_heading(text, level=2)
            except KeyError:
                p = doc.add_paragraph(text)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(14)
        elif line_clean.startswith('### '):
            text = line_clean[4:].strip()
            try:
                doc.add_heading(text, level=3)
            except KeyError:
                p = doc.add_paragraph(text)
                p.runs[0].bold = True
                p.runs[0].font.size = Pt(12)
        elif line_clean.startswith('- ') or line_clean.startswith('* '):
            text = line.lstrip(' -*')
            try:
                p = doc.add_paragraph('', style='List Bullet')
            except KeyError:
                p = doc.add_paragraph('• ')
            _add_formatted_runs(p, text)
        elif re.match(r'^\d+\.\s+', line_clean):
            text = re.sub(r'^\d+\.\s+', '', line_clean)
            try:
                p = doc.add_paragraph('', style='List Number')
            except KeyError:
                p = doc.add_paragraph('• ')
            _add_formatted_runs(p, text)
        else:
            # Normal text with formatting
            p = doc.add_paragraph()
            _add_formatted_runs(p, line_clean)
            
    print(f"Saving to {output_docx_path}...")
    doc.save(output_docx_path)
    print("Done!")

def _add_formatted_runs(paragraph, text):
    """Simple parser to handle inline **bold**"""
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True

if __name__ == "__main__":
    md_to_docx(r"c:\Users\ASUS\Documents\Final Sem\kiosk-app-build\FINAL_PROJECT_REPORT.md",
               r"c:\Users\ASUS\Documents\Final Sem\kiosk-app-build\Workflow\Final Project Report.docx",
               r"c:\Users\ASUS\Documents\Final Sem\kiosk-app-build\Virtual_TryOn_Final_Report.docx")
